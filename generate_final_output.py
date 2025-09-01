#!/usr/bin/env python3
"""
Generate Final Transcription Output Files
Merges all chunk results into final TXT and DOCX files with speaker separation
"""

import json
import os
import logging
from pathlib import Path
from typing import Dict, List, Any, Tuple
from datetime import datetime
import re

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class FinalOutputGenerator:
    """Generates final transcription output files from chunk results"""
    
    def __init__(self, chunk_dir: str = "output/chunk_results", output_dir: str = "output/transcriptions"):
        self.chunk_dir = Path(chunk_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
        # Configuration
        self.chunk_duration = 30
        self.stride_length = 5
        
    def load_all_chunks(self) -> List[Dict[str, Any]]:
        """Load all chunk files and sort them by chunk number"""
        chunk_files = list(self.chunk_dir.glob("chunk_*.json"))
        
        # Sort by chunk number
        chunk_files.sort(key=lambda x: self._extract_chunk_number(x.name))
        
        logger.info(f"Found {len(chunk_files)} chunk files")
        
        chunks = []
        for chunk_file in chunk_files:
            try:
                with open(chunk_file, 'r', encoding='utf-8') as f:
                    chunk_data = json.load(f)
                chunks.append(chunk_data)
            except Exception as e:
                logger.error(f"Error loading {chunk_file}: {e}")
        
        return chunks
    
    def _extract_chunk_number(self, filename: str) -> int:
        """Extract chunk number from filename"""
        try:
            if 'chunk_' in filename:
                parts = filename.split('_')
                if len(parts) >= 2:
                    return int(parts[1])
        except (ValueError, IndexError):
            pass
        return 0
    
    def merge_chunks_with_overlapping(self, chunks: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Merge chunks considering overlapping strategy"""
        if not chunks:
            return {}
        
        # Sort chunks by start time
        sorted_chunks = sorted(chunks, key=lambda x: x.get('transcription_data', {}).get('segments', [{}])[0].get('start', 0))
        
        # Initialize merged data
        merged_data = {
            'total_duration': 0,
            'total_chunks': len(chunks),
            'speakers': {},
            'full_text': '',
            'segments': [],
            'metadata': {
                'generated_at': datetime.now().isoformat(),
                'chunk_count': len(chunks),
                'overlapping_strategy': f"{self.chunk_duration}s chunks with {self.stride_length}s overlap"
            }
        }
        
        # Process each chunk
        for i, chunk in enumerate(sorted_chunks):
            chunk_start = chunk.get('transcription_data', {}).get('segments', [{}])[0].get('start', 0)
            chunk_end = chunk.get('transcription_data', {}).get('segments', [{}])[0].get('end', 0)
            
            # Add chunk duration (accounting for overlap)
            if i == 0:
                merged_data['total_duration'] += chunk_end - chunk_start
            else:
                # For overlapping chunks, only add the non-overlapping portion
                overlap = chunk_end - chunk_start - self.stride_length
                merged_data['total_duration'] += overlap
            
            # Process speaker recognition
            speaker_recognition = chunk.get('speaker_recognition', {})
            if speaker_recognition:
                self._merge_speaker_data(merged_data, speaker_recognition, chunk_start, chunk)
            
            # Add to full text
            chunk_text = chunk.get('text', '')
            if chunk_text:
                if merged_data['full_text']:
                    merged_data['full_text'] += ' ' + chunk_text
                else:
                    merged_data['full_text'] = chunk_text
        
        return merged_data
    
    def _merge_speaker_data(self, merged_data: Dict[str, Any], speaker_recognition: Dict[str, Any], 
                           chunk_start: float, chunk: Dict[str, Any]):
        """Merge speaker data from a chunk"""
        speaker_names = speaker_recognition.get('speaker_names', [])
        speaker_details = speaker_recognition.get('speaker_details', {})
        
        for speaker_name in speaker_names:
            if speaker_name not in merged_data['speakers']:
                merged_data['speakers'][speaker_name] = []
            
            # Get speaker segments
            if speaker_name in speaker_details:
                speaker_data = speaker_details[speaker_name]
                segments = speaker_data.get('segments', [])
                
                for segment in segments:
                    # Convert relative timing to absolute timing
                    absolute_start = chunk_start + segment.get('start', 0)
                    absolute_end = chunk_start + segment.get('end', 0)
                    
                    merged_segment = {
                        'start': absolute_start,
                        'end': absolute_end,
                        'text': segment.get('text', ''),
                        'duration': segment.get('duration', 0),
                        'chunk_source': chunk.get('file_name', 'unknown')
                    }
                    
                    merged_data['speakers'][speaker_name].append(merged_segment)
        
        # Sort segments by start time for each speaker
        for speaker_name in merged_data['speakers']:
            merged_data['speakers'][speaker_name].sort(key=lambda x: x['start'])
    
    def generate_text_output(self, merged_data: Dict[str, Any], output_filename: str = None) -> str:
        """Generate TXT output with speaker separation"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"transcription_{timestamp}.txt"
        
        output_path = self.output_dir / output_filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("TRANSCRIPTION OUTPUT\n")
            f.write("=" * 50 + "\n")
            f.write(f"Generated: {merged_data['metadata']['generated_at']}\n")
            f.write(f"Total chunks: {merged_data['metadata']['chunk_count']}\n")
            f.write(f"Total duration: {merged_data['total_duration']:.2f} seconds\n")
            f.write(f"Overlapping strategy: {merged_data['metadata']['overlapping_strategy']}\n")
            f.write("=" * 50 + "\n\n")
            
            # Write conversation format with chronological speaker turns
            f.write("\n💬 CONVERSATION TRANSCRIPT\n")
            f.write("=" * 50 + "\n\n")
            
            # Collect all segments from all speakers and sort by start time
            all_segments = []
            for speaker_name, segments in merged_data['speakers'].items():
                for segment in segments:
                    all_segments.append({
                        'speaker': speaker_name,
                        'start': segment['start'],
                        'end': segment['end'],
                        'text': segment['text'],
                        'chunk_source': segment['chunk_source']
                    })
            
            # Sort all segments by start time to create chronological conversation
            all_segments.sort(key=lambda x: x['start'])
            
            # Group consecutive segments from the same speaker into conversation turns
            if all_segments:
                current_speaker = None
                current_turn = []
                current_timing = []
                
                for i, segment in enumerate(all_segments):
                    # Check if this should start a new speaker turn
                    should_start_new_turn = False
                    
                    if i == 0:
                        # First segment always starts a new turn
                        should_start_new_turn = True
                    else:
                        # Check if speaker changed or there's a significant gap (>3 seconds)
                        prev_segment = all_segments[i-1]
                        speaker_changed = segment['speaker'] != prev_segment['speaker']
                        time_gap = segment['start'] - prev_segment['end']
                        
                        if speaker_changed or time_gap > 3.0:
                            should_start_new_turn = True
                    
                    if should_start_new_turn and current_turn:
                        # Write the current speaker turn
                        timing_range = f"[{current_timing[0]:.1f}s - {current_timing[-1]:.1f}s]"
                        f.write(f"🎤 {current_speaker} {timing_range}:\n")
                        f.write(f"{' '.join(current_turn)}\n\n")
                        
                        # Start new turn
                        current_speaker = segment['speaker']
                        current_turn = [segment['text']]
                        current_timing = [segment['start'], segment['end']]
                    else:
                        # Continue current turn
                        if current_speaker is None:
                            current_speaker = segment['speaker']
                        current_turn.append(segment['text'])
                        current_timing.append(segment['end'])
                
                # Write the last speaker turn
                if current_turn:
                    timing_range = f"[{current_timing[0]:.1f}s - {current_timing[-1]:.1f}s]"
                    f.write(f"🎤 {current_speaker} {timing_range}:\n")
                    f.write(f"{' '.join(current_turn)}\n\n")
            
            f.write("\n" + "=" * 50 + "\n")
            f.write("SPEAKER SUMMARY\n")
            f.write("=" * 50 + "\n")
            
            # Write speaker summary
            for speaker_name in sorted(merged_data['speakers'].keys()):
                segments = merged_data['speakers'][speaker_name]
                if segments:
                    total_duration = sum(seg['duration'] for seg in segments)
                    segment_count = len(segments)
                    f.write(f"\n🎤 {speaker_name}:\n")
                    f.write(f"  Total segments: {segment_count}\n")
                    f.write(f"  Total duration: {total_duration:.1f}s\n")
                    f.write(f"  Average segment: {total_duration/segment_count:.1f}s\n")
            
            # Write full text
            f.write("\n" + "=" * 50 + "\n")
            f.write("FULL TRANSCRIPTION\n")
            f.write("=" * 50 + "\n")
            f.write(merged_data['full_text'])
        
        logger.info(f"Generated TXT output: {output_path}")
        return str(output_path)
    
    def generate_docx_output(self, merged_data: Dict[str, Any], output_filename: str = None) -> str:
        """Generate DOCX output with speaker separation"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not available. Install with: pip install python-docx")
            return None
        
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"transcription_{timestamp}.docx"
        
        output_path = self.output_dir / output_filename
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('TRANSCRIPTION OUTPUT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f"Generated: {merged_data['metadata']['generated_at']}")
        doc.add_paragraph(f"Total chunks: {merged_data['metadata']['chunk_count']}")
        doc.add_paragraph(f"Total duration: {merged_data['total_duration']:.2f} seconds")
        doc.add_paragraph(f"Overlapping strategy: {merged_data['metadata']['overlapping_strategy']}")
        
        doc.add_page_break()
        
        # Conversation format with chronological speaker turns
        conversation_heading = doc.add_heading('💬 CONVERSATION TRANSCRIPT', level=1)
        conversation_heading.style.font.size = Pt(16)
        conversation_heading.style.font.bold = True
        
        # Collect all segments from all speakers and sort by start time
        all_segments = []
        for speaker_name, segments in merged_data['speakers'].items():
            for segment in segments:
                all_segments.append({
                    'speaker': speaker_name,
                    'start': segment['start'],
                    'end': segment['end'],
                    'text': segment['text'],
                    'chunk_source': segment['chunk_source']
                })
        
        # Sort all segments by start time to create chronological conversation
        all_segments.sort(key=lambda x: x['start'])
        
        # Group consecutive segments from the same speaker into conversation turns
        if all_segments:
            current_speaker = None
            current_turn = []
            current_timing = []
            
            for i, segment in enumerate(all_segments):
                # Check if this should start a new speaker turn
                should_start_new_turn = False
                
                if i == 0:
                    # First segment always starts a new turn
                    should_start_new_turn = True
                else:
                    # Check if speaker changed or there's a significant gap (>3 seconds)
                    prev_segment = all_segments[i-1]
                    speaker_changed = segment['speaker'] != prev_segment['speaker']
                    time_gap = segment['start'] - prev_segment['end']
                    
                    if speaker_changed or time_gap > 3.0:
                        should_start_new_turn = True
                
                if should_start_new_turn and current_turn:
                    # Write the current speaker turn
                    timing_range = f"[{current_timing[0]:.1f}s - {current_timing[-1]:.1f}s]"
                    turn_para = doc.add_paragraph()
                    turn_para.add_run(f"🎤 {current_speaker} ").bold = True
                    turn_para.add_run(timing_range).bold = True
                    turn_para.add_run(":\n")
                    turn_para.add_run(f"{' '.join(current_turn)}")
                    
                    # Start new turn
                    current_speaker = segment['speaker']
                    current_turn = [segment['text']]
                    current_timing = [segment['start'], segment['end']]
                else:
                    # Continue current turn
                    if current_speaker is None:
                        current_speaker = segment['speaker']
                    current_turn.append(segment['text'])
                    current_timing.append(segment['end'])
            
            # Write the last speaker turn
            if current_turn:
                timing_range = f"[{current_timing[0]:.1f}s - {current_timing[-1]:.1f}s]"
                turn_para = doc.add_paragraph()
                turn_para.add_run(f"🎤 {current_speaker} ").bold = True
                turn_para.add_run(timing_range).bold = True
                turn_para.add_run(":\n")
                turn_para.add_run(f"{' '.join(current_turn)}")
        
        doc.add_page_break()
        
        # Speaker summary
        summary_heading = doc.add_heading('SPEAKER SUMMARY', level=1)
        summary_heading.style.font.size = Pt(16)
        summary_heading.style.font.bold = True
        
        for speaker_name in sorted(merged_data['speakers'].keys()):
            segments = merged_data['speakers'][speaker_name]
            if segments:
                total_duration = sum(seg['duration'] for seg in segments)
                segment_count = len(segments)
                doc.add_paragraph(f"🎤 {speaker_name}:")
                doc.add_paragraph(f"  Total segments: {segment_count}")
                doc.add_paragraph(f"  Total duration: {total_duration:.1f}s")
                doc.add_paragraph(f"  Average segment: {total_duration/segment_count:.1f}s")
                doc.add_paragraph()  # Spacing
        
        doc.add_page_break()
        
        # Full transcription
        full_heading = doc.add_heading('FULL TRANSCRIPTION', level=1)
        full_heading.style.font.size = Pt(16)
        full_heading.style.font.bold = True
        
        doc.add_paragraph(merged_data['full_text'])
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Generated DOCX output: {output_path}")
        return str(output_path)
    
    def generate_json_output(self, merged_data: Dict[str, Any], output_filename: str = None) -> str:
        """Generate consolidated JSON output"""
        if not output_filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"transcription_{timestamp}.json"
        
        output_path = self.output_dir / output_filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(merged_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Generated JSON output: {output_path}")
        return str(output_path)
    
    def generate_all_outputs(self) -> Dict[str, str]:
        """Generate all output formats"""
        logger.info("🚀 Starting final output generation...")
        
        # Load and merge chunks
        chunks = self.load_all_chunks()
        if not chunks:
            logger.error("No chunks found to process")
            return {}
        
        # Merge chunks with overlapping strategy
        merged_data = self.merge_chunks_with_overlapping(chunks)
        
        # Generate outputs
        outputs = {}
        
        # TXT output
        txt_path = self.generate_text_output(merged_data)
        if txt_path:
            outputs['txt'] = txt_path
        
        # DOCX output
        docx_path = self.generate_docx_output(merged_data)
        if docx_path:
            outputs['docx'] = docx_path
        
        # JSON output
        json_path = self.generate_json_output(merged_data)
        if json_path:
            outputs['json'] = json_path
        
        logger.info(f"✅ Generated {len(outputs)} output files")
        return outputs


def main():
    """Main entry point"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="Generate final transcription output files from chunk results"
    )
    
    parser.add_argument(
        '--chunk-dir',
        default='output/chunk_results',
        help='Directory containing chunk JSON files'
    )
    
    parser.add_argument(
        '--output-dir',
        default='output/transcriptions',
        help='Directory for output files'
    )
    
    parser.add_argument(
        '--verbose',
        action='store_true',
        help='Enable verbose logging'
    )
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        # Initialize generator
        generator = FinalOutputGenerator(
            chunk_dir=args.chunk_dir,
            output_dir=args.output_dir
        )
        
        # Generate all outputs
        outputs = generator.generate_all_outputs()
        
        if outputs:
            print("\n🎉 OUTPUT GENERATION COMPLETE!")
            print("=" * 50)
            for format_type, file_path in outputs.items():
                print(f"✅ {format_type.upper()}: {file_path}")
        else:
            print("❌ No output files generated")
            return 1
        
        return 0
        
    except Exception as e:
        logger.error(f"Error during output generation: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1


if __name__ == "__main__":
    import sys
    sys.exit(main())
