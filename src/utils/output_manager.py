#!/usr/bin/env python3
"""
Output manager utility
Handles file outputs, logging, and temporary file management
"""

import os
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List
import json

# Check if python-docx is available
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class CustomJSONEncoder(json.JSONEncoder):
    """Custom JSON encoder to handle non-serializable objects"""
    
    def default(self, obj):
        """Handle non-serializable objects by converting them to strings"""
        try:
            # Try to convert to string representation
            return str(obj)
        except:
            # If that fails, return a placeholder
            return f"<Non-serializable object: {type(obj).__name__}>"

class OutputManager:
    """Manages output files, logs, and temporary files"""
    
    def __init__(self, base_output_dir: str = None, run_session_id: str = None):
        """Initialize output manager with base directory"""
        from src.utils.config_manager import config
        
        if base_output_dir is None:
            base_output_dir = config.output.output_dir
        
        self.base_output_dir = Path(base_output_dir)
        self.logs_dir = self.base_output_dir / "logs"
        self.transcriptions_dir = self.base_output_dir / "transcriptions"
        self.temp_dir = self.base_output_dir / "temp"
        
        # Generate run session ID if not provided
        if run_session_id is None:
            self.run_session_id = self._generate_run_session_id()
        else:
            self.run_session_id = run_session_id
        
        # Create directories if they don't exist
        self._ensure_directories()
        
        # Setup logging
        self._setup_logging()
    
    def _generate_run_session_id(self) -> str:
        """Generate a unique run session ID"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"run_{timestamp}"
    
    def get_run_session_dir(self) -> Path:
        """Get the run session directory"""
        return self.transcriptions_dir / self.run_session_id
    
    def _ensure_directories(self):
        """Ensure all output directories exist"""
        self.logs_dir.mkdir(parents=True, exist_ok=True)
        self.transcriptions_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
    
    def _setup_logging(self):
        """Setup logging configuration with log file cleanup"""
        from src.utils.config_manager import config
        
        log_file = self.logs_dir / f"transcription_{datetime.now().strftime('%Y%m%d')}.log"
        
        # Clean log file before each run
        self._clean_log_file(log_file)
        
        # Configure logging with improved timestamp format
        logging.basicConfig(
            level=getattr(logging, config.output.log_level),
            format='%(asctime)s.%(msecs)03d - %(name)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S',
            handlers=[
                logging.FileHandler(log_file, mode='a', encoding='utf-8'),
                logging.StreamHandler()
            ]
        )
        
        self.logger = logging.getLogger('ivrit-ai')
        
        # Log session start
        self.logger.info(f"=== Transcription Session Started: {self.run_session_id} ===")
    
    def _clean_log_file(self, log_file: Path):
        """Clean log file before each run"""
        try:
            if log_file.exists():
                # Create backup of old log
                backup_file = log_file.with_suffix('.log.backup')
                if backup_file.exists():
                    backup_file.unlink()  # Remove old backup
                log_file.rename(backup_file)
                self.logger = logging.getLogger('ivrit-ai')
                self.logger.info(f"Previous log backed up to: {backup_file}")
        except Exception as e:
            # If cleanup fails, just continue with existing log
            pass
    
    def _sanitize_transcription_data(self, data: Any) -> Any:
        """
        Sanitize transcription data to ensure JSON serialization
        
        Args:
            data: The data to sanitize
            
        Returns:
            Sanitized data that can be JSON serialized
        """
        if isinstance(data, dict):
            return {k: self._sanitize_transcription_data(v) for k, v in data.items()}
        elif isinstance(data, list):
            return [self._sanitize_transcription_data(item) for item in data]
        elif isinstance(data, (str, int, float, bool, type(None))):
            return data
        else:
            # Convert non-serializable objects to string
            try:
                return str(data)
            except:
                return f"<Non-serializable: {type(data).__name__}>"
    
    def save_transcription(self, audio_file: str, transcription_data: Any, 
                          model: str = "unknown", engine: str = "unknown") -> str:
        """
        Save transcription results to file
        
        Args:
            audio_file: Original audio file path
            transcription_data: Transcription results
            model: Model used for transcription
            engine: Engine used for transcription
            
        Returns:
            Path to saved transcription file
        """
        # Create run session directory
        run_session_dir = self.get_run_session_dir()
        run_session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create timestamped folder for this transcription session
        audio_name = Path(audio_file).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_folder = f"{timestamp}_{audio_name}"
        session_dir = run_session_dir / session_folder
        session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create filename with sanitized model name
        model_safe = model.replace('/', '_').replace('\\', '_')
        filename = f"transcription_{model_safe}_{engine}.json"
        output_file = session_dir / filename
        
        # Sanitize transcription data to ensure JSON serialization
        sanitized_data = self._sanitize_transcription_data(transcription_data)
        
        # Prepare output data
        output_data = {
            "metadata": {
                "audio_file": audio_file,
                "model": model,
                "engine": engine,
                "timestamp": datetime.now().isoformat(),
                "version": "1.0.0"
            },
            "transcription": sanitized_data
        }
        
        # Save to JSON file with custom encoder
        try:
            with open(output_file, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, ensure_ascii=False, indent=2, cls=CustomJSONEncoder)
            
            self.logger.info(f"Transcription saved to: {output_file}")
            return str(output_file)
        except Exception as e:
            self.logger.error(f"Failed to save transcription to JSON: {e}")
            # Fallback: save as string representation
            fallback_file = output_file.with_suffix('.txt')
            with open(fallback_file, 'w', encoding='utf-8') as f:
                f.write(str(output_data))
            self.logger.info(f"Transcription saved as fallback text to: {fallback_file}")
            return str(fallback_file)
    
    def save_transcription_text(self, audio_file: str, text: str, 
                               model: str = "unknown", engine: str = "unknown") -> str:
        """
        Save transcription as plain text
        
        Args:
            audio_file: Original audio file path
            text: Transcription text
            model: Model used for transcription
            engine: Engine used for transcription
            
        Returns:
            Path to saved text file
        """
        # Create run session directory
        run_session_dir = self.get_run_session_dir()
        run_session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create timestamped folder for this transcription session
        audio_name = Path(audio_file).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_folder = f"{timestamp}_{audio_name}"
        session_dir = run_session_dir / session_folder
        session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create filename with sanitized model name
        model_safe = model.replace('/', '_').replace('\\', '_')
        filename = f"transcription_{model_safe}_{engine}.txt"
        output_file = session_dir / filename
        
        # Prepare text content
        content = f"Transcription of: {audio_file}\n"
        content += f"Model: {model}\n"
        content += f"Engine: {engine}\n"
        content += f"Timestamp: {datetime.now().isoformat()}\n"
        content += "=" * 60 + "\n\n"
        content += text
        
        # Save to text file
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        self.logger.info(f"Transcription text saved to: {output_file}")
        return str(output_file)
    
    def save_transcription_docx(self, audio_file: str, transcription_data: List[Dict[str, Any]], 
                               model: str = "unknown", engine: str = "unknown") -> Optional[str]:
        """
        Save transcription as Word document (.docx) with RTL support and conversation format
        
        Args:
            audio_file: Original audio file path
            transcription_data: List of transcription segments
            model: Model used for transcription
            engine: Engine used for transcription
            
        Returns:
            Path to saved Word document, or None if python-docx not available
        """
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available. Skipping Word document creation.")
            return None
        
        # Create run session directory
        run_session_dir = self.get_run_session_dir()
        run_session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create timestamped folder for this transcription session
        audio_name = Path(audio_file).stem
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        session_folder = f"{timestamp}_{audio_name}"
        session_dir = run_session_dir / session_folder
        session_dir.mkdir(parents=True, exist_ok=True)
        
        # Create filename with sanitized model name
        model_safe = model.replace('/', '_').replace('\\', '_')
        filename = f"transcription_{model_safe}_{engine}.docx"
        output_file = session_dir / filename
        
        # Create Word document
        doc = Document()
        
        # Set document to RTL with comprehensive RTL support
        try:
            section = doc.sections[0]
            # Set RTL direction for the entire document
            bidi_elements = section._sectPr.xpath('./w:bidi')
            if bidi_elements:
                bidi_elements[0].val = '1'  # Enable RTL
            else:
                # Create RTL element if it doesn't exist
                from docx.oxml import OxmlElement
                bidi = OxmlElement('w:bidi')
                bidi.set('w:val', '1')
                section._sectPr.append(bidi)
            
            # Set RTL text flow
            try:
                from docx.oxml import OxmlElement
                textFlow = OxmlElement('w:textFlow')
                textFlow.set('w:val', 'rl-tb')  # Right-to-left, top-to-bottom
                section._sectPr.append(textFlow)
            except Exception as e:
                self.logger.debug(f"Could not set text flow: {e}")
                
        except Exception as e:
            self.logger.warning(f"Could not set RTL direction: {e}. Continuing with default direction.")
        
        # Add title with RTL alignment
        title = doc.add_heading('דוח תמלול', 0)  # Hebrew title: Transcription Report
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set RTL for title paragraph
        for run in title.runs:
            run._element.get_or_add_rPr().get_or_add_rtl()
        
        # Add metadata section with RTL formatting
        metadata_heading = doc.add_heading('מטא-דאטה', level=1)  # Hebrew: Metadata
        metadata_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set RTL for metadata heading
        for run in metadata_heading.runs:
            run._element.get_or_add_rPr().get_or_add_rtl()
        
        metadata_table = doc.add_table(rows=4, cols=2)
        metadata_table.style = 'Table Grid'
        
        metadata_cells = [
            ('קובץ אודיו:', audio_file),
            ('מודל:', model),
            ('מנוע:', engine),
            ('זמן יצירה:', datetime.now().isoformat())
        ]
        
        for i, (key, value) in enumerate(metadata_cells):
            cell_key = metadata_table.cell(i, 0)
            cell_value = metadata_table.cell(i, 1)
            
            cell_key.text = key
            cell_value.text = str(value)
            
            # Set RTL for table cells
            for paragraph in cell_key.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run._element.get_or_add_rPr().get_or_add_rtl()
            
            for paragraph in cell_value.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in paragraph.runs:
                    run._element.get_or_add_rPr().get_or_add_rtl()
        
        doc.add_paragraph()  # Add spacing
        
        # Add transcription content with RTL formatting
        transcript_heading = doc.add_heading('תמלול שיחה', level=1)  # Hebrew: Conversation Transcript
        transcript_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set RTL for transcript heading
        for run in transcript_heading.runs:
            run._element.get_or_add_rPr().get_or_add_rtl()
        
        # Sort segments chronologically
        sorted_segments = sorted(transcription_data, key=lambda x: x.get('start', 0))
        
        # Create conversation format with timestamps and improved punctuation
        for segment in sorted_segments:
            speaker = segment.get('speaker', 'Unknown Speaker')
            text = segment.get('text', '').strip()
            start_time = segment.get('start', 0)
            end_time = segment.get('end', 0)
            
            if text:
                # Improve punctuation for Hebrew text
                text = self._improve_hebrew_punctuation(text)
                
                # Format timestamp
                start_min = int(start_time // 60)
                start_sec = int(start_time % 60)
                end_min = int(end_time // 60)
                end_sec = int(end_time % 60)
                timestamp = f"[{start_min:02d}:{start_sec:02d} - {end_min:02d}:{end_sec:02d}]"
                
                # Create paragraph with RTL alignment
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # RTL alignment
                
                # Set RTL for the entire paragraph using proper method
                try:
                    from docx.oxml import OxmlElement
                    bidi = OxmlElement('w:bidi')
                    bidi.set('w:val', '1')
                    para._element.get_or_add_pPr().append(bidi)
                except Exception as e:
                    self.logger.debug(f"Could not set paragraph RTL: {e}")
                
                # Add timestamp (smaller, gray text)
                timestamp_run = para.add_run(f"{timestamp} ")
                timestamp_run.font.size = Pt(8)
                timestamp_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
                timestamp_run._element.get_or_add_rPr().get_or_add_rtl()
                
                # Add speaker name (bold)
                speaker_run = para.add_run(f"{speaker}: ")
                speaker_run.bold = True
                speaker_run.font.size = Pt(12)
                speaker_run._element.get_or_add_rPr().get_or_add_rtl()
                
                # Add text content with RTL
                text_run = para.add_run(text)
                text_run.font.size = Pt(12)
                text_run._element.get_or_add_rPr().get_or_add_rtl()
                
                # Add spacing between segments
                doc.add_paragraph()
        
        # Save document
        doc.save(str(output_file))
        self.logger.info(f"Transcription Word document saved to: {output_file}")
        return str(output_file)
    
    def _improve_hebrew_punctuation(self, text: str) -> str:
        """
        Improve punctuation for Hebrew text
        
        Args:
            text: Hebrew text to improve
            
        Returns:
            Text with improved punctuation
        """
        if not text:
            return text
        
        # Hebrew punctuation improvements
        improvements = [
            # Fix spacing around Hebrew punctuation
            (r'([א-ת])\s*([,\.!?;:])', r'\1\2'),  # Remove space before punctuation
            (r'([,\.!?;:])\s*([א-ת])', r'\1 \2'),  # Add space after punctuation before Hebrew
            
            # Fix spacing around English punctuation
            (r'([א-ת])\s*([a-zA-Z])', r'\1 \2'),  # Add space between Hebrew and English
            (r'([a-zA-Z])\s*([א-ת])', r'\1 \2'),  # Add space between English and Hebrew
            
            # Fix common Hebrew punctuation patterns
            (r'([א-ת])\s*\.\s*([א-ת])', r'\1. \2'),  # Period between Hebrew words
            (r'([א-ת])\s*,\s*([א-ת])', r'\1, \2'),   # Comma between Hebrew words
            
            # Fix quotation marks for Hebrew
            (r'"([א-ת]+)"', r'"\1"'),  # Ensure proper Hebrew quotation marks
            (r"'([א-ת]+)'", r"'\1'"),   # Ensure proper Hebrew apostrophes
            
            # Fix multiple spaces
            (r'\s+', ' '),  # Replace multiple spaces with single space
            
            # Fix spacing around timestamps and numbers
            (r'([א-ת])\s*(\d+)', r'\1 \2'),  # Space between Hebrew and numbers
            (r'(\d+)\s*([א-ת])', r'\1 \2'),  # Space between numbers and Hebrew
        ]
        
        # Apply improvements
        for pattern, replacement in improvements:
            text = re.sub(pattern, replacement, text)
        
        # Trim whitespace
        text = text.strip()
        
        return text
    
    def get_temp_file(self, prefix: str = "temp", suffix: str = ".tmp") -> str:
        """
        Get a temporary file path
        
        Args:
            prefix: File prefix
            suffix: File suffix
            
        Returns:
            Path to temporary file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
        filename = f"{prefix}_{timestamp}{suffix}"
        return str(self.temp_dir / filename)
    
    def cleanup_temp_files(self, older_than_hours: int = 24):
        """
        Clean up temporary files older than specified hours
        
        Args:
            older_than_hours: Age threshold in hours
        """
        import time
        current_time = time.time()
        threshold = current_time - (older_than_hours * 3600)
        
        cleaned_count = 0
        for temp_file in self.temp_dir.glob("*"):
            if temp_file.is_file() and temp_file.stat().st_mtime < threshold:
                temp_file.unlink()
                cleaned_count += 1
        
        if cleaned_count > 0:
            self.logger.info(f"Cleaned up {cleaned_count} temporary files")
    
    def get_output_stats(self) -> Dict[str, Any]:
        """
        Get statistics about output files
        
        Returns:
            Dictionary with output statistics
        """
        # Count run sessions (parent folders)
        run_sessions = [d for d in self.transcriptions_dir.iterdir() if d.is_dir() and d.name.startswith("run_")]
        
        # Count transcription sessions (subfolders within run sessions)
        transcription_sessions = []
        json_files = []
        txt_files = []
        docx_files = []
        
        for run_session_dir in run_sessions:
            # Count subfolders within this run session
            subfolders = [d for d in run_session_dir.iterdir() if d.is_dir()]
            transcription_sessions.extend(subfolders)
            
            # Count files in all subfolders
            for subfolder in subfolders:
                json_files.extend(subfolder.glob("*.json"))
                txt_files.extend(subfolder.glob("*.txt"))
                docx_files.extend(subfolder.glob("*.docx"))
        
        stats = {
            "logs": {
                "count": len(list(self.logs_dir.glob("*.log"))),
                "size_mb": sum(f.stat().st_size for f in self.logs_dir.glob("*.log")) / (1024 * 1024)
            },
            "transcriptions": {
                "run_sessions": len(run_sessions),
                "sessions": len(transcription_sessions),
                "json_files": len(json_files),
                "txt_files": len(txt_files),
                "docx_files": len(docx_files),
                "total_files": len(json_files) + len(txt_files) + len(docx_files),
                "size_mb": sum(f.stat().st_size for f in json_files + txt_files + docx_files) / (1024 * 1024)
            },
            "temp_files": {
                "count": len(list(self.temp_dir.glob("*"))),
                "size_mb": sum(f.stat().st_size for f in self.temp_dir.glob("*")) / (1024 * 1024)
            }
        }
        
        return stats
    
    def cleanup_all_logs(self, keep_backups: bool = True):
        """Clean all log files, optionally keeping backups"""
        try:
            log_files = list(self.logs_dir.glob("*.log"))
            cleaned_count = 0
            
            for log_file in log_files:
                if keep_backups:
                    # Create backup
                    backup_file = log_file.with_suffix('.log.backup')
                    if backup_file.exists():
                        backup_file.unlink()
                    log_file.rename(backup_file)
                    cleaned_count += 1
                else:
                    # Remove completely
                    log_file.unlink()
                    cleaned_count += 1
            
            if cleaned_count > 0:
                print(f"🧹 Cleaned {cleaned_count} log files")
            else:
                print("✅ No log files to clean")
                
        except Exception as e:
            print(f"⚠️  Warning: Could not clean log files: {e}")
    
    def log_info(self, message: str):
        """Log info message with timestamp"""
        if hasattr(self, 'logger'):
            self.logger.info(f"[INFO] {message}")
        else:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
            print(f"{timestamp} - ivrit-ai - INFO - {message}")
    
    def log_error(self, message: str):
        """Log error message with timestamp"""
        if hasattr(self, 'logger'):
            self.logger.error(f"[ERROR] {message}")
        else:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
            print(f"{timestamp} - ivrit-ai - ERROR - {message}")
    
    def log_warning(self, message: str):
        """Log warning message with timestamp"""
        if hasattr(self, 'logger'):
            self.logger.warning(f"[WARNING] {message}")
        else:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
            print(f"{timestamp} - ivrit-ai - WARNING - {message}")
    
    def log_debug(self, message: str):
        """Log debug message with timestamp"""
        if hasattr(self, 'logger'):
            self.logger.debug(f"[DEBUG] {message}")
        else:
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S.%f')[:-3]
            print(f"{timestamp} - ivrit-ai - DEBUG - {message}") 