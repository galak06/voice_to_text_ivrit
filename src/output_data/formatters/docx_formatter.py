#!/usr/bin/env python3
"""
Enhanced DOCX Formatter
Handles Word document formatting with improved structure, headings, and document splitting
ONLY generates high-quality DOCX from processed text files
"""

import logging
import re
import math
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path  
from datetime import datetime  

from docx import Document  
from docx.shared import Inches, Pt, RGBColor 
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

class DocxFormatter:
    """Enhanced DOCX formatting utilities with RTL support, headings, and document splitting"""
    
    @staticmethod
    def improve_hebrew_punctuation(text: str) -> str:
        """Improve punctuation for Hebrew text"""
        if not text:
            return text
        
        import re
        
        # First, normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Hebrew punctuation improvements
        improvements = [
            # Fix spacing around Hebrew punctuation marks
            (r'([א-ת])\s*([,\.!?;:])', r'\1\2'),  # Remove space before punctuation
            (r'([,\.!?;:])\s*([א-ת])', r'\1 \2'),  # Add space after punctuation before Hebrew
            (r'([א-ת])\s*\.\s*([א-ת])', r'\1. \2'),  # Period between Hebrew words
            (r'([א-ת])\s*,\s*([א-ת])', r'\1, \2'),  # Comma between Hebrew words
            (r'([א-ת])\s*!\s*([א-ת])', r'\1! \2'),  # Exclamation between Hebrew words
            (r'([א-ת])\s*\?\s*([א-ת])', r'\1? \2'),  # Question mark between Hebrew words
            
            # Fix spacing around quotes
            (r'"([א-ת]+)"', r'"\1"'),
            (r"'([א-ת]+)'", r"'\1'"),
            
            # Fix spacing around numbers
            (r'([א-ת])\s*(\d+)', r'\1 \2'),
            (r'(\d+)\s*([א-ת])', r'\1 \2'),
            
            # Fix spacing around English letters
            (r'([א-ת])\s*([a-zA-Z])', r'\1 \2'),
            (r'([a-zA-Z])\s*([א-ת])', r'\1 \2'),
            
            # Fix common Hebrew text issues
            (r'([א-ת])\s*-\s*([א-ת])', r'\1-\2'),  # Hyphens in Hebrew words
            (r'([א-ת])\s*/\s*([א-ת])', r'\1/\2'),  # Slashes in Hebrew words
            
            # Clean up multiple spaces (but preserve word spacing)
            (r'[ \t]+', ' '),
        ]
        
        for pattern, replacement in improvements:
            text = re.sub(pattern, replacement, text)
        
        return text.strip()
    
    @staticmethod
    def create_docx_from_word_ready_text(
        text_file_path: str,
        output_docx_path: str,
        audio_file: str = "merged_chunks",
        model: str = "merged",
        engine: str = "merged",
        split_documents: bool = True,
        chunk_duration_minutes: int = 20
    ) -> bool:
        """
        Create enhanced DOCX file(s) from the Word-ready text file
        
        Args:
            text_file_path: Path to the text file
            output_docx_path: Path for the output DOCX file
            audio_file: Name of the audio file
            model: Model used for transcription
            engine: Engine used for transcription
            split_documents: Whether to split into multiple documents
            chunk_duration_minutes: Duration of each chunk in minutes
        """
        try:
            # Read the Word-ready text file
            with open(text_file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            if not text_content.strip():
                logger.error("Text file is empty or contains only whitespace")
                return False
            
            if split_documents:
                return DocxFormatter._create_split_documents(
                    text_content, output_docx_path, audio_file, model, engine, chunk_duration_minutes
                )
            else:
                return DocxFormatter._create_single_document(
                    text_content, output_docx_path, audio_file, model, engine
                )
            
        except FileNotFoundError:
            logger.error(f"Text file not found: {text_file_path}")
            return False
        except Exception as e:
            logger.error(f"Error creating DOCX from Word-ready text: {e}")
            return False
    
    @staticmethod
    def _create_single_document(
        text_content: str,
        output_docx_path: str,
        audio_file: str,
        model: str,
        engine: str
    ) -> bool:
        """Create a single enhanced DOCX document"""
        try:
            # Create the document
            doc = Document()
            
            # Set document to RTL for Hebrew text
            DocxFormatter._set_document_rtl(doc)
            
            # Add title
            title = doc.add_heading('🎤 TRANSCRIPTION - READY FOR WORD', 0)
            DocxFormatter._set_paragraph_rtl(title)
            
            # Add metadata table
            DocxFormatter._add_metadata_table(doc, audio_file, model, engine)
            
            # Add the enhanced formatted text content
            DocxFormatter._add_enhanced_text_content(doc, text_content)
            
            # Save the document
            doc.save(output_docx_path)
            logger.info(f"✅ Enhanced DOCX created successfully: {output_docx_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error creating single DOCX document: {e}")
            return False
    
    @staticmethod
    def _create_split_documents(
        text_content: str,
        output_docx_path: str,
        audio_file: str,
        model: str,
        engine: str,
        chunk_duration_minutes: int
    ) -> bool:
        """Create multiple DOCX documents split by time chunks"""
        try:
            # Parse the text content to extract segments with timestamps
            segments = DocxFormatter._parse_text_with_timestamps(text_content)
            
            if not segments:
                logger.warning("No segments found, creating single document")
                return DocxFormatter._create_single_document(
                    text_content, output_docx_path, audio_file, model, engine
                )
            
            # Calculate total duration and number of chunks needed
            total_duration = max(segment['end_time'] for segment in segments)
            total_minutes = total_duration / 60
            num_chunks = math.ceil(total_minutes / chunk_duration_minutes)
            
            logger.info(f"📊 Splitting {total_minutes:.1f} minutes into {num_chunks} chunks of {chunk_duration_minutes} minutes each")
            
            # Create output directory if it doesn't exist
            output_dir = Path(output_docx_path).parent
            base_name = Path(output_docx_path).stem
            
            # Split segments into chunks
            chunk_segments = DocxFormatter._split_segments_into_chunks(segments, chunk_duration_minutes)
            
            # Create a document for each chunk
            created_files = []
            for i, chunk_segs in enumerate(chunk_segments, 1):
                chunk_start_time = min(seg['start_time'] for seg in chunk_segs)
                chunk_end_time = max(seg['end_time'] for seg in chunk_segs)
                
                # Create chunk filename
                chunk_filename = f"{base_name}_part_{i:02d}_{chunk_start_time//60:02.0f}min-{chunk_end_time//60:02.0f}min.docx"
                chunk_path = output_dir / chunk_filename
                
                # Create document for this chunk
                doc = Document()
                DocxFormatter._set_document_rtl(doc)
                
                # Add title for this chunk
                title = doc.add_heading(f'🎤 TRANSCRIPTION - PART {i} ({chunk_start_time//60:02.0f}:{chunk_start_time%60:02.0f} - {chunk_end_time//60:02.0f}:{chunk_end_time%60:02.0f})', 0)
                DocxFormatter._set_paragraph_rtl(title)
                
                # Add metadata table
                DocxFormatter._add_metadata_table(doc, audio_file, model, engine, chunk_info=f"Part {i} of {num_chunks}")
                
                # Add content for this chunk
                DocxFormatter._add_enhanced_segments_content(doc, chunk_segs)
                
                # Save the chunk document
                doc.save(str(chunk_path))
                created_files.append(str(chunk_path))
                logger.info(f"✅ Created chunk {i}/{num_chunks}: {chunk_filename}")
            
            # Also create a master document with all content
            master_path = output_dir / f"{base_name}_complete.docx"
            DocxFormatter._create_single_document(text_content, str(master_path), audio_file, model, engine)
            created_files.append(str(master_path))
            
            logger.info(f"✅ Created {len(created_files)} DOCX files total")
            return True
            
        except Exception as e:
            logger.error(f"Error creating split DOCX documents: {e}")
            return False
    
    @staticmethod
    def _parse_text_with_timestamps(text_content: str) -> List[Dict[str, Any]]:
        """Parse text content to extract segments with timestamps"""
        segments = []
        lines = text_content.split('\n')
        
        current_segment = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for timestamp lines
            timestamp_match = re.match(r'\[(\d+):(\d+):(\d+\.?\d*)\s*-\s*(\d+):(\d+):(\d+\.?\d*)\]', line)
            if timestamp_match:
                # Save previous segment if exists
                if current_segment:
                    segments.append(current_segment)
                
                # Start new segment
                start_h, start_m, start_s = timestamp_match.groups()[:3]
                end_h, end_m, end_s = timestamp_match.groups()[3:]
                
                start_time = int(start_h) * 3600 + int(start_m) * 60 + float(start_s)
                end_time = int(end_h) * 3600 + int(end_m) * 60 + float(end_s)
                
                current_segment = {
                    'start_time': start_time,
                    'end_time': end_time,
                    'text': '',
                    'timestamp_line': line
                }
            elif current_segment and not line.startswith('🎤') and not line.startswith('=') and not line.startswith('Generated:') and not line.startswith('Total Chunks:') and not line.startswith('Duration:'):
                # Add text to current segment
                if current_segment['text']:
                    current_segment['text'] += ' ' + line
                else:
                    current_segment['text'] = line
        
        # Add the last segment
        if current_segment:
            segments.append(current_segment)
        
        return segments
    
    @staticmethod
    def _split_segments_into_chunks(segments: List[Dict[str, Any]], chunk_duration_minutes: int) -> List[List[Dict[str, Any]]]:
        """Split segments into time-based chunks"""
        chunk_duration_seconds = chunk_duration_minutes * 60
        chunks = []
        current_chunk = []
        current_chunk_end = 0
        
        for segment in segments:
            # If this segment starts after the current chunk ends, start a new chunk
            if segment['start_time'] >= current_chunk_end:
                if current_chunk:
                    chunks.append(current_chunk)
                current_chunk = [segment]
                current_chunk_end = segment['start_time'] + chunk_duration_seconds
            else:
                # Add to current chunk
                current_chunk.append(segment)
                # Extend chunk end time if this segment goes beyond it
                if segment['end_time'] > current_chunk_end:
                    current_chunk_end = segment['end_time']
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    @staticmethod
    def _add_enhanced_text_content(doc: Document, text_content: str) -> None:
        """Add enhanced formatted text content to the document"""
        try:
            # Parse segments with timestamps
            segments = DocxFormatter._parse_text_with_timestamps(text_content)
            
            if segments:
                # Use enhanced segments content
                DocxFormatter._add_enhanced_segments_content(doc, segments)
            else:
                # Fallback to line-by-line processing
                DocxFormatter._add_enhanced_line_content(doc, text_content)
            
        except Exception as e:
            logger.error(f"Error adding enhanced text content: {e}")
            raise
    
    @staticmethod
    def _add_enhanced_segments_content(doc: Document, segments: List[Dict[str, Any]]) -> None:
        """Add enhanced content from parsed segments"""
        try:
            sentence_count = 0
            
            for i, segment in enumerate(segments):
                # Add timestamp as a heading
                timestamp_heading = doc.add_heading(f"⏰ {segment['timestamp_line']}", level=2)
                DocxFormatter._set_paragraph_rtl(timestamp_heading)
                
                # Process the text content
                text = segment['text'].strip()
                if not text:
                    continue
                
                # Split into sentences and group into logical paragraphs
                sentences = DocxFormatter._split_into_sentences(text)
                paragraphs = DocxFormatter._group_sentences_into_paragraphs(sentences)
                
                for paragraph_sentences in paragraphs:
                    if not paragraph_sentences:
                        continue
                    
                    # Combine sentences in this paragraph
                    paragraph_text = ' '.join(paragraph_sentences)
                    
                    # Add paragraph to document
                    para = doc.add_paragraph(paragraph_text)
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    DocxFormatter._set_paragraph_rtl(para)
                    
                    sentence_count += len(paragraph_sentences)
                    
                    # Add spacing between paragraphs
                    doc.add_paragraph()  # Empty paragraph for spacing
                
                # Add spacing between segments
                if i < len(segments) - 1:
                    doc.add_paragraph()
            
            logger.info(f"✅ Added {len(segments)} segments with {sentence_count} sentences to document")
            
        except Exception as e:
            logger.error(f"Error adding enhanced segments content: {e}")
            raise
    
    @staticmethod
    def _add_enhanced_line_content(doc: Document, text_content: str) -> None:
        """Add enhanced content line by line (fallback method)"""
        try:
            lines = text_content.split('\n')
            sentence_count = 0
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if this is a header line
                if line.startswith('🎤') or line.startswith('=') or line.startswith('Generated:') or line.startswith('Total Chunks:') or line.startswith('Duration:'):
                    if line.startswith('🎤'):
                        header = doc.add_heading(line, level=1)
                        DocxFormatter._set_paragraph_rtl(header)
                    elif line.startswith('='):
                        sep = doc.add_paragraph(line)
                        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        DocxFormatter._set_paragraph_rtl(sep)
                    else:
                        meta = doc.add_paragraph(line)
                        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        DocxFormatter._set_paragraph_rtl(meta)
                
                # Check if this is a timestamp line
                elif line.startswith('[') and ']' in line:
                    timestamp = doc.add_heading(f"⏰ {line}", level=2)
                    DocxFormatter._set_paragraph_rtl(timestamp)
                
                # Check if this is the end section
                elif line.startswith('End of Transcription') or line.startswith('Copy this text'):
                    end = doc.add_paragraph(line)
                    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    DocxFormatter._set_paragraph_rtl(end)
                
                # This is regular text content
                else:
                    # Split into sentences and group into logical paragraphs
                    sentences = DocxFormatter._split_into_sentences(line)
                    paragraphs = DocxFormatter._group_sentences_into_paragraphs(sentences)
                    
                    for paragraph_sentences in paragraphs:
                        if not paragraph_sentences:
                            continue
                        
                        # Combine sentences in this paragraph
                        paragraph_text = ' '.join(paragraph_sentences)
                        
                        # Add paragraph to document
                        text_para = doc.add_paragraph(paragraph_text)
                        text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        DocxFormatter._set_paragraph_rtl(text_para)
                        
                        sentence_count += len(paragraph_sentences)
                        
                        # Add spacing between paragraphs
                        doc.add_paragraph()
            
            logger.info(f"✅ Added {len(lines)} lines with {sentence_count} sentences to document")
            
        except Exception as e:
            logger.error(f"Error adding enhanced line content: {e}")
            raise
    
    @staticmethod
    def _split_into_sentences(text: str) -> List[str]:
        """Split text into sentences for better formatting with improved Hebrew support"""
        if not text:
            return []
        
        # Normalize whitespace and clean text
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Enhanced Hebrew sentence splitting with better punctuation handling
        # Look for sentence boundaries while keeping punctuation with the sentence
        sentence_boundaries = [
            # Standard sentence endings
            r'[.!?]+\s+',
            # Hebrew punctuation marks
            r'[׳״]+\s+',
            # Ellipsis as sentence boundary
            r'\.{3,}\s+',
            # Colon followed by Hebrew text (new topic)
            r':\s+(?=[א-ת])',
            # Semicolon as sentence boundary
            r';\s+',
        ]
        
        # Create pattern that captures the boundary but keeps punctuation with the sentence
        boundary_pattern = '(' + '|'.join(sentence_boundaries) + ')'
        
        # Split while preserving boundaries
        parts = re.split(boundary_pattern, text)
        
        # Reconstruct sentences properly
        sentences = []
        current_sentence = ""
        
        for part in parts:
            if not part or part.isspace():
                continue
            
            # If this part is a boundary pattern, add it to current sentence and finalize
            if re.match(boundary_pattern, part):
                current_sentence += part.rstrip()  # Remove trailing space from boundary
                if current_sentence.strip():
                    sentences.append(current_sentence.strip())
                current_sentence = ""
            else:
                # Regular text part
                current_sentence += part
        
        # Add any remaining text as final sentence
        if current_sentence.strip():
            sentences.append(current_sentence.strip())
        
        # Post-process: clean up and merge fragments
        processed_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Remove standalone punctuation sentences
            if re.match(r'^[.!?׳״;:]+$', sentence):
                continue
            
            # Merge very short fragments with previous sentence if appropriate
            if (len(sentence.split()) < 2 and 
                processed_sentences and 
                not re.search(r'[.!?״׳]$', processed_sentences[-1]) and
                len(processed_sentences[-1].split()) < 10):  # Don't merge to very long sentences
                processed_sentences[-1] += " " + sentence
            else:
                processed_sentences.append(sentence)
        
        return processed_sentences
    
    @staticmethod
    def _group_sentences_into_paragraphs(sentences: List[str]) -> List[List[str]]:
        """Group sentences into logical paragraphs based on content and flow"""
        if not sentences:
            return []
        
        # For very few sentences, group them more intelligently
        if len(sentences) <= 2:
            return [sentences]
        elif len(sentences) <= 4:
            # Split into 2 paragraphs for better readability
            mid = len(sentences) // 2
            return [sentences[:mid], sentences[mid:]]
        
        paragraphs = []
        current_paragraph = []
        
        # Adaptive paragraph sizing based on total sentences
        total_sentences = len(sentences)
        if total_sentences <= 8:
            min_sentences = 2
            max_sentences = 3
        elif total_sentences <= 15:
            min_sentences = 2
            max_sentences = 4
        else:
            min_sentences = 3
            max_sentences = 5
        
        for i, sentence in enumerate(sentences):
            current_paragraph.append(sentence)
            
            # Check if we should end the current paragraph
            should_end_paragraph = False
            
            # Always end if we've reached max sentences
            if len(current_paragraph) >= max_sentences:
                should_end_paragraph = True
            
            # Look ahead: if next sentence starts with topic indicators, end current paragraph
            elif (i < len(sentences) - 1 and 
                  len(current_paragraph) >= min_sentences):
                next_sentence = sentences[i + 1].strip()
                # Topic indicators: question words, transition words, etc.
                topic_indicators = [
                    r'^(אז|כך|לכן|אבל|עכשיו|אחר כך|בנוסף|למעשה|בעיקר|כמו כן)\s',
                    r'^(מה|איך|מתי|איפה|למה|מדוע)\s',
                    r'^(ראשית|שנית|שלישית|לבסוף|לסיכום)\s',
                    r'^(בעבר|כיום|בעתיד|לאחרונה)\s',
                    r'^(תשובה|שאלה|נושא|נקודה)\s*:',  # Q&A patterns
                ]
                
                if any(re.search(pattern, next_sentence, re.IGNORECASE) for pattern in topic_indicators):
                    should_end_paragraph = True
            
            # End paragraph at logical breaking points (long pauses, speaker changes)
            elif ('...' in sentence or '---' in sentence) and len(current_paragraph) >= min_sentences:
                should_end_paragraph = True
            
            # Create balanced paragraph splits - don't let paragraphs get too uneven
            remaining_sentences = len(sentences) - i - 1
            if (remaining_sentences > 0 and 
                  len(current_paragraph) >= min_sentences and
                  remaining_sentences <= max_sentences):  # Remaining sentences can form a good paragraph
                # Check if current sentence is a good ending point
                if re.search(r'[.!?״׳]$', sentence.strip()):
                    should_end_paragraph = True
            
            # Force paragraph break if this is the last sentence
            if i == len(sentences) - 1:
                should_end_paragraph = True
            
            if should_end_paragraph:
                if current_paragraph:
                    paragraphs.append(current_paragraph[:])  # Copy the list
                current_paragraph = []
        
        # Handle any remaining sentences
        if current_paragraph:
            if paragraphs and len(current_paragraph) <= 2 and len(paragraphs[-1]) < max_sentences:
                # Merge short remainder with last paragraph if it's not too long
                paragraphs[-1].extend(current_paragraph)
            else:
                paragraphs.append(current_paragraph)
        
        # Post-process: ensure no empty paragraphs and optimize paragraph sizes
        paragraphs = [p for p in paragraphs if p and any(s.strip() for s in p)]
        
        # Merge very short paragraphs with adjacent ones for better balance
        optimized_paragraphs = []
        for paragraph in paragraphs:
            if (len(paragraph) == 1 and 
                optimized_paragraphs and 
                len(optimized_paragraphs[-1]) < max_sentences):
                # Merge single sentence with previous paragraph
                optimized_paragraphs[-1].extend(paragraph)
            else:
                optimized_paragraphs.append(paragraph)
        
        return optimized_paragraphs if optimized_paragraphs else paragraphs
    
    @staticmethod
    def _add_word_ready_text_content(doc: Document, text_content: str) -> None:
        """Add the Word-ready text content to the document with proper formatting"""
        try:
            # Split content into lines
            lines = text_content.split('\n')
            
            # Process each line
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # Check if this is a header line
                if line.startswith('🎤') or line.startswith('=') or line.startswith('Generated:') or line.startswith('Total Chunks:') or line.startswith('Duration:'):
                    # Add header as a formatted paragraph
                    if line.startswith('🎤'):
                        header = doc.add_heading(line, level=1)
                        DocxFormatter._set_paragraph_rtl(header)
                    elif line.startswith('='):
                        # Add separator line
                        sep = doc.add_paragraph(line)
                        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        DocxFormatter._set_paragraph_rtl(sep)
                    else:
                        # Add metadata line
                        meta = doc.add_paragraph(line)
                        meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        DocxFormatter._set_paragraph_rtl(meta)
                
                # Check if this is a timestamp line
                elif line.startswith('[') and ']' in line:
                    # Add timestamp as a formatted paragraph
                    timestamp = doc.add_paragraph(line)
                    timestamp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    # Make timestamp bold
                    for run in timestamp.runs:
                        run.bold = True
                        run.font.size = Pt(11)
                    DocxFormatter._set_paragraph_rtl(timestamp)
                
                # Check if this is the end section
                elif line.startswith('End of Transcription') or line.startswith('Copy this text'):
                    # Add end section
                    end = doc.add_paragraph(line)
                    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    DocxFormatter._set_paragraph_rtl(end)
                
                # This is regular text content
                else:
                    # Add regular text content
                    text_para = doc.add_paragraph(line)
                    text_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right-aligned for Hebrew
                    DocxFormatter._set_paragraph_rtl(text_para)
            
            logger.info(f"✅ Added {len(lines)} lines of Word-ready text content to document")
            
        except Exception as e:
            logger.error(f"Error adding Word-ready text content: {e}")
            raise
    
    @staticmethod
    def _set_document_rtl(doc: Document):
        """Set document to RTL direction with simplified Hebrew support"""
        try:
            # Simplified RTL setup - just set basic document properties
            # This avoids the complex RTL setup that can cause opening issues
            
            # Set page margins for Hebrew text (right margin smaller for RTL)
            section = doc.sections[0]
            section.left_margin = Inches(1.5)   # Larger left margin for Hebrew text
            section.right_margin = Inches(0.5)  # Smaller right margin
            
            logger.info("Document RTL settings applied successfully")
        except Exception as e:
            logger.warning(f"Could not set RTL direction: {e}")
    
    @staticmethod
    def _set_paragraph_rtl(paragraph):
        """Set paragraph to RTL alignment with simplified Hebrew text support"""
        try:
            # Simplified RTL - just set right alignment for Hebrew text
            # This avoids the complex RTL direction properties that can cause issues
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception as e:
            logger.warning(f"Could not set paragraph RTL: {e}")
    
    @staticmethod
    def _set_text_run_rtl(text_run):
        """Set text run to RTL direction for Hebrew text - simplified approach"""
        try:
            # Simplified RTL - just ensure proper font handling for Hebrew
            # Don't set complex RTL properties that can cause opening issues
            pass
        except Exception as e:
            logger.debug(f"Could not set text run RTL: {e}")
    
    @staticmethod
    def _add_metadata_table(doc: Document, audio_file: str, model: str, engine: str, chunk_info: str = None):
        """Add enhanced metadata table to document"""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add headers
        header_cells = table.rows[0].cells
        header_cells[0].text = 'מטא-דאטה'
        header_cells[1].text = 'ערך'
        
        # Set RTL for header cells
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                DocxFormatter._set_paragraph_rtl(paragraph)
        
        # Add data rows
        data = [
            ('קובץ שמע', Path(audio_file).name),
            ('מודל', model),
            ('מנוע', engine),
            ('תאריך', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        # Add chunk info if provided
        if chunk_info:
            data.append(('חלק', chunk_info))
        
        for key, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)
            
            # Set RTL for all cells
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    DocxFormatter._set_paragraph_rtl(paragraph) 