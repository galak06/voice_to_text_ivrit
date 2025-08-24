#!/usr/bin/env python3
"""
Create DOCX from existing JSON transcription
"""

import os
import json
import logging
import time
import re
from pathlib import Path

# Set up simple logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

def clean_text_for_xml(text):
    """Clean text to make it XML compatible"""
    if not text:
        return text
    
    # Remove control characters except for common whitespace
    # Keep only: tab (9), newline (10), carriage return (13), and normal characters (32-126 + Hebrew range)
    cleaned = ""
    for char in text:
        code = ord(char)
        if (code == 9 or code == 10 or code == 13 or  # Tab, newline, carriage return
            (32 <= code <= 126) or  # ASCII printable
            (0x0590 <= code <= 0x05FF) or  # Hebrew block
            (0xFB1D <= code <= 0xFB4F)):  # Hebrew presentation forms
            cleaned += char
        else:
            # Replace problematic characters with space
            cleaned += " "
    
    # Normalize whitespace
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned

def create_docx_from_json():
    """Create DOCX from existing JSON file"""
    
    # Find the latest JSON file
    transcriptions_dir = "/Users/gilcohen/voic_to_text_docker/output/transcriptions"
    
    # Look for the most recent directory
    latest_dir = None
    latest_time = 0
    
    for item in os.listdir(transcriptions_dir):
        item_path = os.path.join(transcriptions_dir, item)
        if os.path.isdir(item_path):
            # Look for subdirectories with timestamps
            for subitem in os.listdir(item_path):
                subitem_path = os.path.join(item_path, subitem)
                if os.path.isdir(subitem_path):
                    # Get modification time
                    mod_time = os.path.getmtime(subitem_path)
                    if mod_time > latest_time:
                        latest_time = mod_time
                        latest_dir = subitem_path
    
    if not latest_dir:
        logger.error("No transcription directory found")
        return False
    
    # Find JSON file in the latest directory
    json_file = None
    for file in os.listdir(latest_dir):
        if file.endswith('_transcription.json'):
            json_file = os.path.join(latest_dir, file)
            break
    
    if not json_file:
        logger.error("No JSON transcription file found")
        return False
    
    logger.info(f"Found JSON file: {json_file}")
    
    # Load JSON data
    with open(json_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create DOCX
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        docx_filename = json_file.replace('_transcription.json', '_transcription.docx')
        logger.info(f"Creating DOCX: {docx_filename}")
        
        # Create a new document
        doc = Document()
        
        # Set RTL support for Hebrew
        def set_paragraph_rtl(paragraph):
            """Set paragraph direction to RTL for Hebrew text"""
            try:
                pPr = paragraph._element.get_or_add_pPr()
                bidi = pPr.find('.//w:bidi', pPr.nsmap)
                if bidi is None:
                    from docx.oxml import OxmlElement
                    bidi = OxmlElement('w:bidi')
                    pPr.append(bidi)
            except Exception:
                pass  # Fallback if RTL setting fails
        
        # Add title
        title = doc.add_heading('תמלול שיחה', level=1)
        set_paragraph_rtl(title)
        
        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"קובץ שמע: {os.path.basename(data.get('audio_file', ''))}\n")
        metadata_para.add_run(f"מודל: {data.get('model_name', '')}\n")
        metadata_para.add_run(f"מנוע: {data.get('engine', '')}\n")
        metadata_para.add_run(f"סך הכל מילים: {data.get('total_words', 0)}\n")
        metadata_para.add_run(f"מספר דוברים: {data.get('speaker_count', 0)}\n")
        metadata_para.add_run(f"חלקים שוחזרו: {data.get('merged_chunks', 0)}\n")
        metadata_para.add_run(f"חלקים שדולגו: {data.get('skipped_chunks', 0)}\n")
        set_paragraph_rtl(metadata_para)
        
        # Add separator
        doc.add_paragraph("=" * 50)
        
        # Add transcription content
        content_heading = doc.add_heading('תוכן התמלול', level=2)
        set_paragraph_rtl(content_heading)
        
        # Add each speaker segment
        speakers_data = data.get('speakers', {})
        for speaker_name, segments in speakers_data.items():
            speaker_heading = doc.add_heading(f'🎤 {speaker_name}', level=3)
            set_paragraph_rtl(speaker_heading)
            
            for segment in segments:
                # Create paragraph for each segment
                segment_para = doc.add_paragraph()
                
                # Add timestamp
                start_time = segment.get('start', 0)
                end_time = segment.get('end', 0)
                timestamp_text = f"[{start_time:.1f}s - {end_time:.1f}s]: "
                timestamp_run = segment_para.add_run(clean_text_for_xml(timestamp_text))
                timestamp_run.font.size = Pt(10)
                timestamp_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
                
                # Add text content (clean it first)
                original_text = segment.get('text', '')
                cleaned_text = clean_text_for_xml(original_text)
                text_run = segment_para.add_run(cleaned_text)
                text_run.font.size = Pt(12)
                
                set_paragraph_rtl(segment_para)
                
                # Add some spacing between segments
                doc.add_paragraph()
        
        # Save the document
        doc.save(docx_filename)
        logger.info(f"✅ DOCX created successfully: {docx_filename}")
        return True
        
    except ImportError:
        logger.warning("⚠️  python-docx not available - cannot create DOCX")
        return False
    except Exception as e:
        logger.error(f"❌ Failed to create DOCX: {e}")
        import traceback
        logger.error(f"Error details: {traceback.format_exc()}")
        return False

def main():
    """Main function"""
    logger.info("🚀 Creating DOCX from JSON transcription...")
    
    if create_docx_from_json():
        logger.info("🎉 DOCX creation completed successfully!")
        return 0
    else:
        logger.error("❌ DOCX creation failed")
        return 1

if __name__ == "__main__":
    import sys
    sys.exit(main())
