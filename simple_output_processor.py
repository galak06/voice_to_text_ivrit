#!/usr/bin/env python3
"""
Simple script to process quarterly meeting transcription output from chunks
"""

import os
import json
import glob
from pathlib import Path
from datetime import datetime

def load_chunk_data():
    """Load all chunk files and extract text"""
    chunk_files = glob.glob("output/chunk_results/chunk_*.json")
    chunk_files.sort()
    
    print(f"📁 Found {len(chunk_files)} chunk files")
    
    segments = []
    combined_text = ""
    
    for chunk_file in chunk_files:
        try:
            with open(chunk_file, 'r', encoding='utf-8') as f:
                chunk_data = json.load(f)
            
            if chunk_data.get('status') == 'completed' and chunk_data.get('text'):
                segment = {
                    'start': chunk_data.get('start_time', 0.0),
                    'end': chunk_data.get('end_time', 30.0),
                    'text': chunk_data.get('text', ''),
                    'words': chunk_data.get('words_estimated', 0)
                }
                segments.append(segment)
                combined_text += chunk_data.get('text', '') + ' '
                
        except Exception as e:
            print(f"⚠️ Error loading {chunk_file}: {e}")
    
    return segments, combined_text.strip()

def deduplicate_overlapping_text(segments):
    """Simple deduplication of overlapping text"""
    if not segments:
        return segments
    
    deduplicated = [segments[0]]  # Keep first segment
    
    for i in range(1, len(segments)):
        current = segments[i]
        previous = deduplicated[-1]
        
        # Check for overlap
        if current['start'] < previous['end']:
            # Calculate overlap duration
            overlap_duration = previous['end'] - current['start']
            overlap_ratio = overlap_duration / (current['end'] - current['start'])
            
            # If significant overlap, remove overlapping text from current segment
            if overlap_ratio > 0.1:  # More than 10% overlap
                # Simple approach: remove first part of current text
                words = current['text'].split()
                words_to_remove = int(len(words) * overlap_ratio * 0.5)  # Remove half of overlap
                if words_to_remove > 0:
                    current['text'] = ' '.join(words[words_to_remove:])
                    current['start'] = previous['end']  # Adjust start time
        
        deduplicated.append(current)
    
    return deduplicated

def create_output_files(segments, combined_text):
    """Create output files"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = f"output/transcriptions/run_{timestamp}"
    os.makedirs(output_dir, exist_ok=True)
    
    # Create JSON output
    json_data = {
        "text": combined_text,
        "full_text": combined_text,
        "segments": segments,
        "metadata": {
            "total_segments": len(segments),
            "total_characters": len(combined_text),
            "total_words": len(combined_text.split()),
            "audio_duration": segments[-1]['end'] if segments else 0,
            "model": "ivrit-ai/whisper-large-v3-ct2",
            "engine": "ctranslate2-whisper",
            "language": "he",
            "generated_at": datetime.now().isoformat(),
            "source_file": "examples/audio/voice/פגישה רביעית.wav"
        }
    }
    
    json_file = f"{output_dir}/transcription_quarterly_meeting.json"
    with open(json_file, 'w', encoding='utf-8') as f:
        json.dump(json_data, f, ensure_ascii=False, indent=2)
    
    # Create text output
    text_file = f"{output_dir}/transcription_quarterly_meeting.txt"
    with open(text_file, 'w', encoding='utf-8') as f:
        f.write("פגישה רבעית - תמלול\n")
        f.write("=" * 50 + "\n\n")
        f.write(combined_text)
    
    # Create simple DOCX (if python-docx is available)
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('פגישה רבעית - תמלול', 0)
        
        # Add metadata
        doc.add_paragraph(f"תאריך: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"משך הקלטה: {segments[-1]['end']/60:.1f} דקות" if segments else "לא ידוע")
        doc.add_paragraph(f"סך מילים: {len(combined_text.split())}")
        doc.add_paragraph("")
        
        # Add main text
        doc.add_paragraph(combined_text)
        
        # Set RTL for Hebrew
        for paragraph in doc.paragraphs:
            paragraph.alignment = 1  # Right alignment
        
        docx_file = f"{output_dir}/transcription_quarterly_meeting.docx"
        doc.save(docx_file)
        
        print(f"✅ Created DOCX: {docx_file}")
        
    except ImportError:
        print("⚠️ python-docx not available, skipping DOCX creation")
    
    return json_file, text_file, output_dir

def main():
    """Main function"""
    print("🎤 Processing Quarterly Meeting Transcription")
    print("=" * 60)
    
    # Load chunk data
    segments, combined_text = load_chunk_data()
    
    if not segments:
        print("❌ No segments found!")
        return
    
    print(f"✅ Loaded {len(segments)} segments with {len(combined_text)} characters")
    
    # Deduplicate overlapping text
    print("🔄 Deduplicating overlapping text...")
    deduplicated_segments = deduplicate_overlapping_text(segments)
    deduplicated_text = ' '.join([seg['text'] for seg in deduplicated_segments])
    
    print(f"✅ After deduplication: {len(deduplicated_segments)} segments, {len(deduplicated_text)} characters")
    
    # Create output files
    print("📝 Creating output files...")
    json_file, text_file, output_dir = create_output_files(deduplicated_segments, deduplicated_text)
    
    print(f"\n✅ Output processing completed!")
    print(f"📁 Output directory: {output_dir}")
    print(f"📄 JSON file: {json_file}")
    print(f"📄 Text file: {text_file}")
    
    # Show file sizes
    json_size = os.path.getsize(json_file)
    text_size = os.path.getsize(text_file)
    print(f"📊 JSON file size: {json_size:,} bytes")
    print(f"📊 Text file size: {text_size:,} bytes")
    
    # Show sample text
    print(f"\n📝 Sample text (first 200 characters):")
    print("-" * 50)
    print(deduplicated_text[:200] + "..." if len(deduplicated_text) > 200 else deduplicated_text)

if __name__ == "__main__":
    main()
