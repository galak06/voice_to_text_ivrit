#!/usr/bin/env python3
"""
Create DOCX from consolidated transcription JSON
"""
import json
import sys
import os
import glob
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_docx_from_json(json_file_path, output_path=None):
    """Create a DOCX file from JSON transcription data"""
    
    try:
        # Load JSON data
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        # Create document
        doc = Document()
        
        # Set document properties
        doc.core_properties.author = "Voice-to-Text Transcription System"
        doc.core_properties.title = "Hebrew Audio Transcription"
        doc.core_properties.created = datetime.now()
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "×ª××œ×•×œ ×©×™×—×” ×‘×¢×‘×¨×™×ª - Hebrew Audio Transcription"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add title
        title = doc.add_heading('×ª××œ×•×œ ××•×“×™×• ×‘×¢×‘×¨×™×ª', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata section
        doc.add_heading('×¤×¨×˜×™ ×”×ª××œ×•×œ', level=1)
        
        metadata = data.get('metadata', {})
        meta_para = doc.add_paragraph()
        meta_para.add_run('×ª××¨×™×š: ').bold = True
        meta_para.add_run(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
        
        meta_para.add_run('××©×š ×”×§×œ×˜×”: ').bold = True
        duration_minutes = metadata.get('total_duration', 0) / 60
        meta_para.add_run(f"{duration_minutes:.1f} ×“×§×•×ª\n")
        
        meta_para.add_run('××¡×¤×¨ ××™×œ×™×: ').bold = True
        meta_para.add_run(f"{metadata.get('total_words', 0):,}\n")
        
        meta_para.add_run('××¡×¤×¨ ×ª×•×•×™×: ').bold = True
        meta_para.add_run(f"{metadata.get('total_characters', 0):,}\n")
        
        meta_para.add_run('××¡×¤×¨ ×§×˜×¢×™×: ').bold = True
        meta_para.add_run(f"{metadata.get('total_chunks', 0)}")
        
        # Add transcription content
        doc.add_heading('×ª×•×›×Ÿ ×”×ª××œ×•×œ', level=1)
        
        # Get the full text
        full_text = data.get('full_text', '')
        
        if full_text:
            # Split text into paragraphs (roughly every 5-6 sentences)
            sentences = full_text.split('. ')
            paragraphs = []
            current_paragraph = []
            
            for i, sentence in enumerate(sentences):
                sentence = sentence.strip()
                if sentence:
                    # Add period back if it's not the last sentence
                    if i < len(sentences) - 1:
                        sentence += '.'
                    current_paragraph.append(sentence)
                    
                    # Create new paragraph every 4-5 sentences
                    if len(current_paragraph) >= 4:
                        paragraphs.append(' '.join(current_paragraph))
                        current_paragraph = []
            
            # Add remaining sentences
            if current_paragraph:
                paragraphs.append(' '.join(current_paragraph))
            
            # Add each paragraph to document
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph(para_text.strip())
                    # Set RTL direction for Hebrew text
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Set font and formatting
                    for run in para.runs:
                        run.font.name = 'David'
                        run.font.size = Pt(12)
        else:
            doc.add_paragraph("××™×Ÿ ×ª×•×›×Ÿ ×–××™×Ÿ ×‘×§×•×‘×¥ JSON")
        
        # Add footer with timestamp chunks
        if 'chunks' in data and data['chunks']:
            doc.add_page_break()
            doc.add_heading('×¤×™×¨×•×˜ ×–×× ×™×', level=1)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '×§×˜×¢'
            hdr_cells[1].text = '×–××Ÿ ×”×ª×—×œ×”'
            hdr_cells[2].text = '×–××Ÿ ×¡×™×•×'
            hdr_cells[3].text = '××¡×¤×¨ ××™×œ×™×'
            
            # Add sample of first 10 chunks
            for chunk in data['chunks'][:10]:
                row_cells = table.add_row().cells
                row_cells[0].text = str(chunk.get('chunk_number', ''))
                row_cells[1].text = f"{chunk.get('start_time', 0):.1f}s"
                row_cells[2].text = f"{chunk.get('end_time', 0):.1f}s"
                row_cells[3].text = str(chunk.get('word_count', 0))
            
            if len(data['chunks']) > 10:
                para = doc.add_paragraph(f"... ×•×”××©×š ×¢×“ ×§×˜×¢ {len(data['chunks'])}")
        
        # Save document
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"output/transcriptions/transcription_{timestamp}.docx"
        
        doc.save(output_path)
        print(f"âœ… DOCX created successfully: {output_path}")
        return output_path
        
    except Exception as e:
        print(f"âŒ Error creating DOCX: {e}")
        return None

def main():
    if len(sys.argv) > 1:
        json_file = sys.argv[1]
    else:
        # Find the most recent JSON file in timestamped folders
        pattern = "output/transcriptions/run_*/×¤×’×™×©×”*_transcription.json"
        json_files = glob.glob(pattern)
        if not json_files:
            pattern = "output/transcriptions/run_*/*_transcription.json"
            json_files = glob.glob(pattern)
        if not json_files:
            pattern = "output/transcriptions/transcription_*.json"
            json_files = glob.glob(pattern)

        if not json_files:
            print("âŒ No JSON transcription files found")
            return 1

        # Get the most recent file
        json_file = max(json_files, key=os.path.getctime)
        print(f"ğŸ“ Using most recent JSON file: {json_file}")

    if not os.path.exists(json_file):
        print(f"âŒ JSON file not found: {json_file}")
        return 1

    # Create DOCX in the same folder as the JSON file
    json_dir = os.path.dirname(json_file)
    json_basename = os.path.splitext(os.path.basename(json_file))[0]
    docx_output = os.path.join(json_dir, f"{json_basename}.docx")

    result = create_docx_from_json(json_file, docx_output)
    return 0 if result else 1

if __name__ == "__main__":
    sys.exit(main())